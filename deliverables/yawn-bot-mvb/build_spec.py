from __future__ import annotations

import math
import re
import shutil
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parents[1]
ASSETS = ROOT / "assets"
BUILD = WORKSPACE / "build" / "yawn-bot-mvb"
SOURCE_MD = ROOT / "YAWN.bot MVB-1 Engineering Specification.md"
SOURCE_RENDER = Path(r"C:\Users\dave\Downloads\yawn_bot_render_front.png")
OUTPUT_DOCX = ROOT / "YAWN.bot MVB-1 Engineering Specification.docx"

INK = "17212B"
SLATE = "405261"
BLUE = "1F4E79"
PALE = "E8EEF5"
CYAN = "17E7FF"
MAGENTA = "FF3CCF"
LIME = "A9FF68"
YELLOW = "E8FF03"
ORANGE = "FF7300"
RED = "FF5C6C"
BLACK = "05070A"
WHITE = "FFFFFF"
GRAY = "71808D"
LIGHT = "F5F7FA"


def rgb(hex_value: str) -> tuple[int, int, int]:
    value = hex_value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


FONT_REGULAR = Path(r"C:\Windows\Fonts\segoeui.ttf")
FONT_SEMIBOLD = Path(r"C:\Windows\Fonts\seguisb.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\segoeuib.ttf")
FONT_MONO = Path(r"C:\Windows\Fonts\consola.ttf")


def font(size: int, weight: str = "regular") -> ImageFont.FreeTypeFont:
    choices = {
        "regular": FONT_REGULAR,
        "semibold": FONT_SEMIBOLD,
        "bold": FONT_BOLD,
        "mono": FONT_MONO,
    }
    return ImageFont.truetype(str(choices[weight]), size=size)


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        if draw.textlength(test, font=fnt) <= width or not current:
            current = test
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def multiline(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    fnt: ImageFont.FreeTypeFont,
    fill: tuple[int, int, int],
    width: int,
    spacing: int = 8,
    anchor: str | None = None,
) -> int:
    x, y = xy
    lines = wrap(draw, text, fnt, width)
    line_height = fnt.getbbox("Ag")[3] - fnt.getbbox("Ag")[1]
    for line in lines:
        if anchor == "mm":
            draw.text((x, y + line_height / 2), line, font=fnt, fill=fill, anchor="mm")
        else:
            draw.text((x, y), line, font=fnt, fill=fill)
        y += line_height + spacing
    return y


def base_canvas(title: str, subtitle: str, width: int = 1800, height: int = 1100):
    image = Image.new("RGB", (width, height), rgb(WHITE))
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 92), fill=rgb(BLACK))
    draw.text((70, 28), "YAWN.bot", font=font(38, "bold"), fill=rgb(CYAN))
    draw.text((width - 70, 45), "MVB-1 • ENGINEERING SPEC", font=font(22, "semibold"), fill=rgb(WHITE), anchor="rm")
    draw.text((70, 126), title, font=font(48, "bold"), fill=rgb(INK))
    draw.text((70, 188), subtitle, font=font(24), fill=rgb(SLATE))
    draw.line((70, 232, width - 70, 232), fill=rgb(CYAN), width=5)
    return image, draw


def rounded(draw: ImageDraw.ImageDraw, box, fill_color=WHITE, outline=PALE, radius=24, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=rgb(fill_color), outline=rgb(outline), width=width)


def label(draw: ImageDraw.ImageDraw, xy, text: str, color: str = BLUE):
    x, y = xy
    fnt = font(18, "bold")
    tw = draw.textlength(text, font=fnt)
    draw.rounded_rectangle((x, y, x + tw + 28, y + 34), radius=8, fill=rgb(color))
    draw.text((x + 14, y + 6), text, font=fnt, fill=rgb(WHITE))


def arrow(draw: ImageDraw.ImageDraw, start, end, color=SLATE, width=5, head=16):
    draw.line((start, end), fill=rgb(color), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    for offset in (math.pi * 0.82, -math.pi * 0.82):
        p = (end[0] + head * math.cos(angle + offset), end[1] + head * math.sin(angle + offset))
        draw.line((end, p), fill=rgb(color), width=width)


def double_arrow(draw: ImageDraw.ImageDraw, start, end, color=BLUE, width=4, head=13):
    arrow(draw, start, end, color=color, width=width, head=head)
    arrow(draw, end, start, color=color, width=width, head=head)


def save(image: Image.Image, name: str):
    ASSETS.mkdir(parents=True, exist_ok=True)
    image.save(ASSETS / name, format="PNG", optimize=True, dpi=(180, 180))


def prepare_render():
    image = Image.open(SOURCE_RENDER).convert("RGB")
    # Retain the atmosphere while removing excess top/side void.
    crop = image.crop((180, 145, 1185, 1940))
    crop.save(ASSETS / "render-design-intent.png", format="PNG", optimize=True, dpi=(180, 180))


def product_family():
    image, draw = base_canvas(
        "Modular product family",
        "One identity, two compute classes, and optional backs — each with an honest boundary",
    )
    cards = [
        (80, 285, 460, 930, "CARD 0", "QR / NFC identity", "85.60 × 53.98 mm", "TRUE WALLET", CYAN, "0"),
        (500, 285, 880, 930, "MVB-1", "Browser Screen Core", "93 × 57 × 18 mm", "CHROMIUM", LIME, "1"),
        (920, 285, 1300, 930, "BACKS", "Battery + sensors", "22–34 mm total", "MODULAR", MAGENTA, "+"),
        (1340, 285, 1720, 930, "SQUARE", "CoreS3-Lite native", "≈54 × 54 × 16.5 mm", "NO BROWSER", ORANGE, "S"),
    ]
    for x1, y1, x2, y2, name, subtitle, dims, boundary, accent, glyph in cards:
        rounded(draw, (x1, y1, x2, y2), fill_color=LIGHT, outline=accent, radius=28, width=5)
        draw.text((x1 + 30, y1 + 28), name, font=font(29, "bold"), fill=rgb(INK))
        draw.text((x1 + 30, y1 + 70), subtitle, font=font(21), fill=rgb(SLATE))
        if name == "CARD 0":
            device = (x1 + 72, y1 + 180, x2 - 72, y1 + 365)
        elif name == "SQUARE":
            device = (x1 + 90, y1 + 155, x2 - 90, y1 + 405)
        else:
            device = (x1 + 102, y1 + 130, x2 - 102, y1 + 440)
        draw.rounded_rectangle(device, radius=20, fill=rgb(BLACK), outline=rgb(accent), width=8)
        dx1, dy1, dx2, dy2 = device
        if name == "CARD 0":
            qx, qy = dx2 - 108, dy1 + 38
            for yy in range(5):
                for xx in range(5):
                    if (xx * 3 + yy * 5 + xx * yy) % 4 != 0:
                        draw.rectangle((qx + xx * 13, qy + yy * 13, qx + xx * 13 + 9, qy + yy * 13 + 9), fill=rgb(WHITE))
            draw.text((dx1 + 24, dy1 + 34), "yawn.bot/dave", font=font(18, "mono"), fill=rgb(CYAN))
        else:
            draw.text(((dx1 + dx2) / 2, (dy1 + dy2) / 2), glyph, font=font(105, "bold"), fill=rgb(accent), anchor="mm")
            draw.line((dx1 + 15, dy1 + 15, dx2 - 15, dy1 + 15), fill=rgb(accent), width=4)
        draw.text(((x1 + x2) / 2, y1 + 495), dims, font=font(23, "semibold"), fill=rgb(INK), anchor="mm")
        label(draw, (x1 + 30, y1 + 555), boundary, accent)
        if name == "CARD 0":
            note = "Public identity only. No secret is encoded. Opens any phone browser."
        elif name == "MVB-1":
            note = "Primary reference: Pi Zero 2 W + HDMI capacitive display."
        elif name == "BACKS":
            note = "Battery, I2S mic, service-removable CSI camera, and controls."
        else:
            note = "Small integrated client implemented in LVGL / ESP-IDF."
        multiline(draw, (x1 + 30, y1 + 604), note, font(19), rgb(SLATE), x2 - x1 - 60, spacing=5)
    draw.text((900, 1025), "Build behavior first. Reduce thickness only after measured evidence.", font=font(25, "semibold"), fill=rgb(BLUE), anchor="mm")
    save(image, "01-product-family.png")


def dimensioned_views():
    image, draw = base_canvas(
        "MVB-1 target envelope",
        "Primary 3.5-inch browser build — dimensions in millimeters",
    )
    label(draw, (75, 260), "TARGET", CYAN)
    # Front view, portrait.
    x, y, w, h = 200, 330, 310, 505
    draw.rounded_rectangle((x, y, x + w, y + h), radius=16, fill=rgb(BLACK), outline=rgb(INK), width=5)
    draw.rounded_rectangle((x + 13, y + 13, x + w - 13, y + h - 13), radius=10, outline=rgb(CYAN), width=6)
    draw.text((x + w / 2, y + h / 2 - 20), "480 × 800", font=font(31, "bold"), fill=rgb(LIME), anchor="mm")
    draw.text((x + w / 2, y + h / 2 + 27), "CAPACITIVE", font=font(19, "mono"), fill=rgb(CYAN), anchor="mm")
    # Side buttons.
    draw.rounded_rectangle((x + w + 3, y + 180, x + w + 18, y + 238), radius=5, fill=rgb(MAGENTA))
    draw.rounded_rectangle((x + w + 3, y + 275, x + w + 18, y + 333), radius=5, fill=rgb(LIME))
    # Width dimension.
    draw.line((x, y + h + 28, x, y + h + 90), fill=rgb(GRAY), width=3)
    draw.line((x + w, y + h + 28, x + w, y + h + 90), fill=rgb(GRAY), width=3)
    double_arrow(draw, (x, y + h + 66), (x + w, y + h + 66))
    draw.text((x + w / 2, y + h + 94), "57.0 TARGET", font=font(22, "bold"), fill=rgb(BLUE), anchor="mm")
    # Height dimension.
    draw.line((x - 82, y, x - 20, y), fill=rgb(GRAY), width=3)
    draw.line((x - 82, y + h, x - 20, y + h), fill=rgb(GRAY), width=3)
    double_arrow(draw, (x - 55, y), (x - 55, y + h))
    draw.text((x - 106, y + h / 2), "93.0 TARGET", font=font(22, "bold"), fill=rgb(BLUE), anchor="mm")
    draw.text((x + w / 2, y - 32), "FRONT", font=font(22, "semibold"), fill=rgb(SLATE), anchor="mm")

    # Side view.
    sx, sy, sw, sh = 670, 330, 98, 505
    draw.rounded_rectangle((sx, sy, sx + sw, sy + sh), radius=10, fill=rgb(LIGHT), outline=rgb(INK), width=5)
    draw.rectangle((sx + 11, sy + 16, sx + 31, sy + sh - 16), fill=rgb(CYAN))
    draw.rectangle((sx + 35, sy + 40, sx + 67, sy + sh - 40), fill=rgb(SLATE))
    draw.text((sx + sw / 2, sy - 32), "SIDE", font=font(22, "semibold"), fill=rgb(SLATE), anchor="mm")
    draw.line((sx, sy + sh + 28, sx, sy + sh + 90), fill=rgb(GRAY), width=3)
    draw.line((sx + sw, sy + sh + 28, sx + sw, sy + sh + 90), fill=rgb(GRAY), width=3)
    double_arrow(draw, (sx, sy + sh + 66), (sx + sw, sy + sh + 66))
    draw.text((sx + sw / 2, sy + sh + 94), "18.0 TARGET", font=font(22, "bold"), fill=rgb(BLUE), anchor="mm")

    # Back view with modules.
    bx, by, bw, bh = 950, 330, 310, 505
    draw.rounded_rectangle((bx, by, bx + bw, by + bh), radius=16, fill=rgb(LIGHT), outline=rgb(INK), width=5)
    draw.rounded_rectangle((bx + 55, by + 70, bx + bw - 55, by + 315), radius=16, fill=rgb(PALE), outline=rgb(BLUE), width=4)
    draw.text((bx + bw / 2, by + 155), "PI ZERO 2 W", font=font(23, "bold"), fill=rgb(BLUE), anchor="mm")
    draw.text((bx + bw / 2, by + 197), "65 × 30", font=font(20), fill=rgb(SLATE), anchor="mm")
    draw.rounded_rectangle((bx + 115, by + 345, bx + 195, by + 425), radius=10, fill=rgb(BLACK), outline=rgb(MAGENTA), width=5)
    draw.ellipse((bx + 140, by + 370, bx + 170, by + 400), fill=rgb(MAGENTA))
    draw.text((bx + bw / 2, by - 32), "BACK / POD", font=font(22, "semibold"), fill=rgb(SLATE), anchor="mm")
    label(draw, (bx + 28, by + 443), "EST. LENS ≤34", MAGENTA)

    # Notes panel.
    nx1, ny1, nx2, ny2 = 1340, 300, 1725, 885
    rounded(draw, (nx1, ny1, nx2, ny2), fill_color=LIGHT, outline=PALE, radius=24, width=4)
    draw.text((nx1 + 28, ny1 + 28), "DATUM & FIT NOTES", font=font(25, "bold"), fill=rgb(INK))
    notes = [
        ("VERIFIED", "Display module: 88.87 × 52.56 × 7.15", LIME),
        ("TARGET", "Outer envelope: 93 × 57 × 18", CYAN),
        ("ESTIMATE", "Battery total: 22–27 thick", ORANGE),
        ("ESTIMATE", "Camera lens zone: ≤34 thick", MAGENTA),
        ("RULE", "Screen glass is datum A; left and bottom edges are B/C", BLUE),
        ("GATE", "Model the purchased HDMI and USB cable ends before final CAD", RED),
    ]
    yy = ny1 + 88
    for tag, note, color in notes:
        label(draw, (nx1 + 28, yy), tag, color)
        yy += 48
        yy = multiline(draw, (nx1 + 28, yy), note, font(20), rgb(SLATE), nx2 - nx1 - 56, spacing=5) + 22
    draw.text((900, 1030), "Not to scale for manufacturing • replace TARGET/ESTIMATE values with prototype measurements", font=font(22, "semibold"), fill=rgb(RED), anchor="mm")
    save(image, "02-dimensioned-views.png")


def exploded_stack():
    image, draw = base_canvas(
        "Exploded enclosure stack",
        "Screw-serviceable modular construction — front at upper left",
    )
    layers = [
        ("01", "FRONT BEZEL", "PETG / ASA • 1.4–1.6 wall", CYAN, 0),
        ("02", "DISPLAY", "Waveshare 24037 • supplier part", BLUE, 85),
        ("03", "0.3 GASKET", "Poron • no glass point loads", LIME, 170),
        ("04", "MIDFRAME", "Pi, cable folds, buttons, mic duct", CYAN, 255),
        ("05", "PI ZERO 2 W", "M2.5 standoffs • antenna keepout", BLUE, 340),
        ("06", "SERVICE BACK", "M2 screws + heat-set inserts", SLATE, 425),
        ("07", "OPTIONAL BACK", "PiSugar / sensor pod • deeper shell", MAGENTA, 510),
    ]
    base_x, base_y = 180, 305
    for idx, (num, name, note, accent, offset) in enumerate(layers):
        x = base_x + offset
        y = base_y + offset * 0.37
        w, h = 620, 330
        fill = BLACK if idx in (0, 1) else LIGHT
        outline = accent
        draw.rounded_rectangle((x, y, x + w, y + h), radius=26, fill=rgb(fill), outline=rgb(outline), width=7)
        draw.line((x + 35, y + 40, x + w - 35, y + 40), fill=rgb(accent), width=5)
        text_color = WHITE if fill == BLACK else INK
        draw.text((x + 38, y + 73), num, font=font(30, "bold"), fill=rgb(accent))
        draw.text((x + 105, y + 70), name, font=font(28, "bold"), fill=rgb(text_color))
        draw.text((x + 105, y + 115), note, font=font(20), fill=rgb(WHITE if fill == BLACK else SLATE))
        if name == "DISPLAY":
            draw.rounded_rectangle((x + 90, y + 160, x + w - 90, y + h - 28), radius=12, outline=rgb(CYAN), width=6)
        if name == "PI ZERO 2 W":
            draw.rounded_rectangle((x + 145, y + 165, x + w - 145, y + h - 25), radius=12, fill=rgb(PALE), outline=rgb(BLUE), width=5)
        if name == "OPTIONAL BACK":
            draw.ellipse((x + w - 155, y + 160, x + w - 55, y + 260), fill=rgb(BLACK), outline=rgb(MAGENTA), width=6)
    # Side notes.
    nx1, ny1, nx2, ny2 = 1270, 300, 1720, 910
    rounded(draw, (nx1, ny1, nx2, ny2), fill_color=WHITE, outline=PALE, radius=24, width=4)
    draw.text((nx1 + 30, ny1 + 30), "NON-NEGOTIABLE", font=font(27, "bold"), fill=rgb(INK))
    bullets = [
        "No screw, boss, or rib over a pouch cell.",
        "Power down before servicing the CSI ribbon.",
        "Keep metal and battery foil away from the Pi antenna end.",
        "Use a manual camera shutter and a visible lock flag.",
        "Publish native CAD, STEP, STL, 3MF, and the exact hardware revision.",
    ]
    yy = ny1 + 100
    for item in bullets:
        draw.ellipse((nx1 + 32, yy + 8, nx1 + 48, yy + 24), fill=rgb(CYAN))
        yy = multiline(draw, (nx1 + 64, yy), item, font(21), rgb(SLATE), nx2 - nx1 - 100, spacing=5) + 26
    label(draw, (nx1 + 30, ny2 - 75), "SERVICEABLE ≠ HOT-SWAP", RED)
    save(image, "03-exploded-stack.png")


def electrical_architecture():
    image, draw = base_canvas(
        "Electrical architecture",
        "Proposed pin allocation for the 3.5-inch HDMI build",
    )
    # Pi center.
    px1, py1, px2, py2 = 650, 350, 1135, 770
    rounded(draw, (px1, py1, px2, py2), fill_color=BLACK, outline=CYAN, radius=28, width=7)
    draw.text(((px1 + px2) / 2, py1 + 65), "RASPBERRY PI ZERO 2 W", font=font(34, "bold"), fill=rgb(WHITE), anchor="mm")
    draw.text(((px1 + px2) / 2, py1 + 115), "Linux + Chromium kiosk", font=font(23), fill=rgb(CYAN), anchor="mm")
    ports = ["mini HDMI", "USB OTG", "CSI", "GPIO / I2S", "I2C", "5 V rail"]
    yy = py1 + 175
    for p in ports:
        draw.rounded_rectangle((px1 + 85, yy, px2 - 85, yy + 42), radius=10, fill=rgb(INK), outline=rgb(SLATE), width=2)
        draw.text(((px1 + px2) / 2, yy + 21), p, font=font(19, "mono"), fill=rgb(WHITE), anchor="mm")
        yy += 50

    blocks = [
        (80, 310, 470, 500, "5 V INPUT", "2.5–3 A regulated\nmeasure undervoltage", ORANGE),
        (80, 600, 470, 800, "PISUGAR 3", "1200 mAh nominal\noptional Stage 5", MAGENTA),
        (1300, 285, 1710, 475, "DISPLAY D1", "480 × 800 HDMI\nUSB touch + power", CYAN),
        (1300, 520, 1710, 700, "CAMERA", "CSI Module 3\nservice-removable", MAGENTA),
        (1300, 745, 1710, 940, "MIC + CONTROLS", "I2S GPIO18/19/20\nbuttons GPIO5/6/13", LIME),
    ]
    for x1, y1, x2, y2, title, note, accent in blocks:
        rounded(draw, (x1, y1, x2, y2), fill_color=LIGHT, outline=accent, radius=24, width=5)
        draw.text((x1 + 28, y1 + 30), title, font=font(27, "bold"), fill=rgb(INK))
        draw.multiline_text((x1 + 28, y1 + 80), note, font=font(21), fill=rgb(SLATE), spacing=8)
    arrow(draw, (470, 405), (650, 470), ORANGE, width=7)
    draw.text((550, 395), "5 V", font=font(18, "bold"), fill=rgb(ORANGE), anchor="mm")
    arrow(draw, (470, 695), (650, 670), MAGENTA, width=7)
    draw.text((550, 720), "5 V + I2C", font=font(18, "bold"), fill=rgb(MAGENTA), anchor="mm")
    arrow(draw, (1135, 420), (1300, 375), CYAN, width=7)
    draw.text((1210, 385), "HDMI + USB", font=font(18, "bold"), fill=rgb(CYAN), anchor="mm")
    arrow(draw, (1135, 570), (1300, 610), MAGENTA, width=7)
    draw.text((1215, 565), "CSI", font=font(18, "bold"), fill=rgb(MAGENTA), anchor="mm")
    arrow(draw, (1135, 715), (1300, 825), LIME, width=7)
    draw.text((1215, 755), "GPIO", font=font(18, "bold"), fill=rgb(LIME), anchor="mm")
    label(draw, (78, 930), "GATE: BOOT + CAPTURE + UPLOAD WITHOUT UNDERVOLTAGE", RED)
    save(image, "04-electrical-architecture.png")


def software_architecture():
    image, draw = base_canvas(
        "Software and repository boundary",
        "Keep protocol, device client, operating-system bridge, and hosted runtime separable",
    )
    rows = [
        (285, [
            (90, 510, "yawn-ai/.yawn", "MIT protocol / schemas", CYAN),
            (560, 980, "yawn-bot device client", "Next.js route + state machine", LIME),
            (1030, 1450, "yawn-hardware", "CAD / BOM / validation", MAGENTA),
            (1500, 1730, "web-game", "hosted runtime • private", ORANGE),
        ]),
        (520, [
            (170, 650, "PAIRING API", "one-time request • scoped cookie", BLUE),
            (720, 1190, "DEVICE ROUTE", "480 × 800 • PWA shell • media", CYAN),
            (1260, 1680, "MEDIA INTAKE", "auth • validate • idempotency", MAGENTA),
        ]),
        (765, [
            (250, 710, "ANOTTERKIOSK / PI OS", "current Chromium • recovery", SLATE),
            (780, 1240, "yawn-device-bridge", "libgpiod → uinput F13/F14/F15", LIME),
            (1310, 1640, "ALSA / V4L2", "mic + camera device gates", ORANGE),
        ]),
    ]
    centers = []
    for y, blocks in rows:
        row_centers = []
        for x1, x2, title, note, accent in blocks:
            rounded(draw, (x1, y, x2, y + 150), fill_color=LIGHT, outline=accent, radius=22, width=5)
            draw.text(((x1 + x2) / 2, y + 44), title, font=font(24, "bold"), fill=rgb(INK), anchor="mm")
            draw.text(((x1 + x2) / 2, y + 94), note, font=font(18), fill=rgb(SLATE), anchor="mm")
            row_centers.append(((x1 + x2) / 2, y + 75))
        centers.append(row_centers)
    # Connections.
    arrow(draw, (300, 435), (405, 520), CYAN, width=5)
    arrow(draw, (770, 435), (955, 520), LIME, width=5)
    arrow(draw, (1240, 435), (1470, 520), MAGENTA, width=5)
    arrow(draw, (1615, 435), (1470, 520), ORANGE, width=5)
    arrow(draw, (955, 670), (1010, 765), CYAN, width=5)
    arrow(draw, (405, 670), (480, 765), BLUE, width=5)
    arrow(draw, (1470, 670), (1475, 765), MAGENTA, width=5)
    # Browser/device silhouette.
    draw.rounded_rectangle((70, 970, 1730, 1045), radius=18, fill=rgb(BLACK), outline=rgb(CYAN), width=5)
    draw.text((900, 1008), "PHYSICAL DEVICE: display + touch + GPIO buttons + microphone + service camera", font=font(23, "semibold"), fill=rgb(WHITE), anchor="mm")
    draw.text((900, 1080), "QR contains a public URL or one-time pairing request — never a long-lived account credential", font=font(21, "bold"), fill=rgb(RED), anchor="mm")
    save(image, "05-software-architecture.png")


def interaction_state_machine():
    image, draw = base_canvas(
        "Interaction and privacy state machine",
        "Touch, keyboard, and GPIO events share one transition model",
    )
    states = [
        (110, 430, 350, 610, "LOCKED", "no media tracks", BLACK),
        (420, 430, 660, 610, "READY", "paired + online", CYAN),
        (730, 430, 970, 610, "ARMED", "hold ≥300 ms", BLUE),
        (1040, 330, 1310, 500, "MIC ACTIVE", "lime rail", LIME),
        (1040, 550, 1310, 720, "CAM ACTIVE", "magenta rail", MAGENTA),
        (1390, 430, 1640, 610, "REVIEW", "discard or send", YELLOW),
    ]
    for x1, y1, x2, y2, title, note, accent in states:
        fill = BLACK if title == "LOCKED" else LIGHT
        rounded(draw, (x1, y1, x2, y2), fill_color=fill, outline=accent, radius=28, width=7)
        draw.text(((x1 + x2) / 2, y1 + 60), title, font=font(28, "bold"), fill=rgb(WHITE if fill == BLACK else INK), anchor="mm")
        draw.text(((x1 + x2) / 2, y1 + 112), note, font=font(20), fill=rgb(CYAN if fill == BLACK else SLATE), anchor="mm")
    arrow(draw, (350, 520), (420, 520), CYAN, width=6)
    arrow(draw, (660, 520), (730, 520), BLUE, width=6)
    arrow(draw, (970, 485), (1040, 420), LIME, width=6)
    arrow(draw, (970, 555), (1040, 630), MAGENTA, width=6)
    arrow(draw, (1310, 420), (1390, 490), YELLOW, width=6)
    arrow(draw, (1310, 630), (1390, 560), YELLOW, width=6)
    # Return path.
    draw.line((1515, 610, 1515, 820, 540, 820, 540, 610), fill=rgb(CYAN), width=5)
    arrow(draw, (540, 690), (540, 610), CYAN, width=5)
    draw.text((1030, 850), "discard / sent / recover", font=font(21, "semibold"), fill=rgb(BLUE), anchor="mm")
    # Lock override.
    draw.line((1180, 720, 1180, 905, 230, 905, 230, 610), fill=rgb(RED), width=6)
    arrow(draw, (230, 690), (230, 610), RED, width=6)
    label(draw, (615, 875), "LOCK / PAGE HIDE / CRASH → STOP TRACKS + CLEAR BLOBS", RED)
    # Input legend.
    inputs = [("TOUCH", CYAN), ("F13 MIC", LIME), ("F14 CAMERA", MAGENTA), ("F15 LOCK", RED)]
    xx = 285
    for text, color in inputs:
        label(draw, (xx, 270), text, color)
        xx += 300
    draw.text((900, 1020), "Release is a first-class event. A stuck key or lost process must fail closed.", font=font(24, "bold"), fill=rgb(INK), anchor="mm")
    save(image, "06-interaction-state-machine.png")


def thickness_comparison():
    image, draw = base_canvas(
        "Thickness reality and product path",
        "Millimeters — off-the-shelf validation precedes custom electronics",
    )
    items = [
        ("ISO CARD", 0.76, "reference", GRAY),
        ("CUSTOM GOAL", 9.0, "future estimate", CYAN),
        ("CORES3-LITE", 16.5, "native client", ORANGE),
        ("MVB-1 CORE", 18.0, "target", LIME),
        ("+ BATTERY", 25.5, "estimate", BLUE),
        ("+ LENS ZONE", 34.0, "estimate", MAGENTA),
    ]
    x0, y_base = 170, 900
    max_height = 560
    scale = max_height / 34.0
    bar_w, gap = 180, 90
    for i, (name, value, qualifier, color) in enumerate(items):
        x = x0 + i * (bar_w + gap)
        height = max(10, value * scale)
        draw.rounded_rectangle((x, y_base - height, x + bar_w, y_base), radius=16, fill=rgb(color), outline=rgb(INK), width=3)
        draw.text((x + bar_w / 2, y_base - height - 45), f"{value:g} mm", font=font(28, "bold"), fill=rgb(INK), anchor="mm")
        draw.text((x + bar_w / 2, y_base + 36), name, font=font(22, "bold"), fill=rgb(INK), anchor="mm")
        draw.text((x + bar_w / 2, y_base + 73), qualifier, font=font(18), fill=rgb(SLATE), anchor="mm")
    draw.line((110, y_base, 1720, y_base), fill=rgb(INK), width=4)
    rounded(draw, (90, 270, 640, 420), fill_color=LIGHT, outline=CYAN, radius=20, width=4)
    draw.text((120, 300), "OFF-THE-SHELF PURPOSE", font=font(25, "bold"), fill=rgb(INK))
    multiline(draw, (120, 345), "Prove interaction, pairing, control placement, service, thermal behavior, and demand.", font(20), rgb(SLATE), 490, spacing=5)
    rounded(draw, (1130, 270, 1710, 420), fill_color=LIGHT, outline=MAGENTA, radius=20, width=4)
    draw.text((1160, 300), "CUSTOM-HARDWARE PURPOSE", font=font(25, "bold"), fill=rgb(INK))
    multiline(draw, (1160, 345), "Remove HDMI/USB cable loops and integrate display, charging, sensors, and privacy power domains.", font(20), rgb(SLATE), 520, spacing=5)
    save(image, "07-thickness-comparison.png")


def generate_figures():
    ASSETS.mkdir(parents=True, exist_ok=True)
    prepare_render()
    product_family()
    dimensioned_views()
    exploded_stack()
    electrical_architecture()
    software_architecture()
    interaction_state_machine()
    thickness_comparison()


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips=9360, indent=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\*\*.*?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\[[^\]]+\]\[[^\]]+\]|https?://[^\s|]+)"
)


def add_inline(paragraph, text: str, refs: dict[str, str]):
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(*rgb(BLUE))
        elif token.startswith("[") and "](" in token:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if m:
                add_hyperlink(paragraph, m.group(1), m.group(2))
        elif token.startswith("[") and "][" in token:
            m = re.match(r"\[([^\]]+)\]\[([^\]]+)\]", token)
            if m:
                url = refs.get(m.group(2).upper())
                if url:
                    add_hyperlink(paragraph, m.group(1), url)
                else:
                    paragraph.add_run(m.group(1))
        elif token.startswith("http"):
            url = token.rstrip(".,;)")
            suffix = token[len(url) :]
            add_hyperlink(paragraph, url, url)
            if suffix:
                paragraph.add_run(suffix)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def paragraph_border(paragraph, color=CYAN, size=18, space=1, side="left"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*rgb(GRAY))
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(*rgb(INK))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    style_specs = {
        "Title": (30, BLACK, 0, 14),
        "Subtitle": (15, SLATE, 0, 14),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, SLATE, 10, 5),
    }
    for name, (size, color, before, after) in style_specs.items():
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(*rgb(color))
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    for name, size, color, italic in (
        ("Figure Caption", 9, SLATE, True),
        ("Table Text", 9, INK, False),
        ("Source Text", 8, SLATE, False),
        ("Callout", 11, INK, False),
        ("Code Block", 8.5, INK, False),
    ):
        if name not in [s.name for s in doc.styles]:
            s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            s = doc.styles[name]
        s.font.name = "Consolas" if name == "Code Block" else "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(*rgb(color))
        s.font.italic = italic
        s.paragraph_format.space_after = Pt(5)
        if name == "Figure Caption":
            s.paragraph_format.keep_with_next = False
        if name == "Callout":
            s.paragraph_format.left_indent = Inches(0.18)
            s.paragraph_format.right_indent = Inches(0.10)
            s.paragraph_format.space_before = Pt(5)
            s.paragraph_format.space_after = Pt(8)
        if name == "Code Block":
            s.paragraph_format.left_indent = Inches(0.22)
            s.paragraph_format.space_before = Pt(4)
            s.paragraph_format.space_after = Pt(8)
            s.paragraph_format.line_spacing = 1.0

    if "YB List" not in [s.name for s in doc.styles]:
        s = doc.styles.add_style("YB List", WD_STYLE_TYPE.PARAGRAPH)
    else:
        s = doc.styles["YB List"]
    for name in ("YB List",):
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(11)
        s.paragraph_format.left_indent = Inches(0.28)
        s.paragraph_format.first_line_indent = Inches(-0.18)
        s.paragraph_format.space_after = Pt(3)
        s.paragraph_format.line_spacing = 1.15


def add_header_footer(section):
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    table.columns[0].width = Inches(3.3)
    table.columns[1].width = Inches(3.2)
    left, right = table.rows[0].cells
    set_cell_shading(left, BLACK)
    set_cell_shading(right, BLACK)
    for cell in (left, right):
        set_cell_margins(cell, top=40, bottom=40, start=90, end=90)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("YAWN.bot")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(*rgb(CYAN))
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("YB-MVB1-ES-001 • REV 1.0")
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(*rgb(WHITE))
    footer = section.footer
    p = footer.paragraphs[0]
    p.add_run("BUILDABLE REFERENCE • TARGETS REQUIRE PHYSICAL VALIDATION • ")
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(7.5)
    p.runs[0].font.color.rgb = RGBColor(*rgb(GRAY))
    add_page_number(p)


def set_image_alt_text(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("YAWN.bot")
    r.font.name = "Consolas"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*rgb(CYAN))
    paragraph_border(p, CYAN, size=28, side="bottom")

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("MVB-1 Engineering Specification")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("A buildable, modular, open-source browser-card reference")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(ASSETS / "render-design-intent.png"), width=Inches(2.65))
    set_image_alt_text(shape, "YAWN.bot product rendering", "Black portrait screen with cyan perimeter status line and two side controls; visual intent only.")

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    set_table_width(table, width_twips=7700, indent=800)
    values = [
        ("DOCUMENT", "YB-MVB1-ES-001 • Revision 1.0"),
        ("DATE / STATUS", "11 July 2026 • prototype release"),
        ("PRIMARY BUILD", "Pi Zero 2 W + 3.5-inch 480 × 800 HDMI touch"),
    ]
    for row, (key, value) in zip(table.rows, values):
        set_row_cant_split(row)
        row.cells[0].width = Inches(1.45)
        row.cells[1].width = Inches(3.90)
        set_cell_shading(row.cells[0], BLACK)
        set_cell_shading(row.cells[1], LIGHT)
        for cell in row.cells:
            set_cell_margins(cell, top=65, bottom=65, start=95, end=95)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        rr = p0.add_run(key)
        rr.bold = True
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor(*rgb(CYAN))
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        rr = p1.add_run(value)
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(*rgb(INK))

    p = doc.add_paragraph(style="Callout")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Build the first real unit as a Linux browser terminal. Keep the wallet-thin identity layer and square native client as separate modules.").bold = True
    paragraph_border(p, MAGENTA, size=24, side="left")
    doc.add_page_break()


def add_contents(doc: Document, headings: list[str]):
    p = doc.add_paragraph("Document map", style="Heading 1")
    p.paragraph_format.page_break_before = False
    p = doc.add_paragraph()
    p.add_run("Use this specification as the purchase boundary, build sequence, and verification record. Dimensions marked TARGET or ESTIMATE are not yet prototype measurements.")
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    for idx, text in enumerate(headings, 1):
        row = table.add_row()
        if idx == 1:
            set_repeat_table_header(row)
        set_row_cant_split(row)
        row.cells[0].width = Inches(0.48)
        row.cells[1].width = Inches(5.92)
        for cell in row.cells:
            set_cell_margins(cell, top=55, bottom=55, start=95, end=95)
        set_cell_shading(row.cells[0], BLACK)
        if idx % 2 == 0:
            set_cell_shading(row.cells[1], LIGHT)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r = p0.add_run(str(idx).zfill(2))
        r.bold = True
        r.font.color.rgb = RGBColor(*rgb(CYAN))
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.add_run(text).bold = True
    p = doc.add_paragraph(style="Callout")
    p.add_run("Release rule: no public claim for wallet thickness, runtime, camera readiness, offline behavior, or security without the matching validation record.")
    paragraph_border(p, RED, size=22, side="left")
    doc.add_page_break()


def add_numbering_definition(doc: Document, ordered: bool, start: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"A1B2{abstract_id:04X}"[-8:])
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), f"C3D4{abstract_id:04X}"[-8:])
    abstract.append(tmpl)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_el = OxmlElement("w:start")
    start_el.set(qn("w:val"), str(start))
    lvl.append(start_el)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "\uf0b7")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "480")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "480")
    ind.set(qn("w:hanging"), "240")
    p_pr.append(ind)
    lvl.append(p_pr)
    if not ordered:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Symbol")
        r_fonts.set(qn("w:hAnsi"), "Symbol")
        r_fonts.set(qn("w:hint"), "default")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    # OOXML requires all abstractNum elements to precede every concrete num.
    # Appending abstractNum after an existing num makes Word repair and merge lists.
    existing_nums = numbering.findall(qn("w:num"))
    if existing_nums:
        numbering.insert(list(numbering).index(existing_nums[0]), abstract)
    else:
        numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def parse_references(lines: list[str]) -> dict[str, str]:
    refs: dict[str, str] = {}
    for line in lines:
        m = re.match(r"^\[([^\]]+)\]:\s+(https?://\S+)", line.strip())
        if m:
            refs[m.group(1).upper()] = m.group(2)
    return refs


def split_table_row(line: str) -> list[str]:
    text = line.strip().strip("|")
    return [cell.strip() for cell in text.split("|")]


def add_markdown_table(doc: Document, rows: list[list[str]], refs: dict[str, str]):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_width(table)
    # A compact but explicit width heuristic.
    max_lens = [max((len(row[c]) if c < len(row) else 0) for row in rows) for c in range(cols)]
    weights = [max(6, min(42, value)) for value in max_lens]
    total = sum(weights)
    widths = [9360 * w / total for w in weights]
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        set_row_cant_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[c_idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, PALE)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            p.paragraph_format.space_after = Pt(0)
            value = values[c_idx] if c_idx < len(values) else ""
            add_inline(p, value, refs)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(*rgb(BLUE))


def add_figure(doc: Document, image_path: Path, alt: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    width = Inches(6.35)
    shape = p.add_run().add_picture(str(image_path), width=width)
    set_image_alt_text(shape, image_path.stem.replace("-", " "), alt)


def parse_body(doc: Document, lines: list[str], refs: dict[str, str]):
    i = 0
    in_code = False
    code_lines: list[str] = []
    numbered_id: int | None = None
    bullet_id: int | None = None
    page_break_sections = {1, 3, 4, 5, 6, 7, 8, 9, 10, 11, 13, 14, 17}
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if re.match(r"^\[[^\]]+\]:\s+https?://", stripped):
            i += 1
            continue
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph(style="Code Block")
                p.add_run("\n".join(code_lines))
                paragraph_border(p, BLUE, size=12, side="left")
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped or stripped == "---":
            numbered_id = None
            bullet_id = None
            i += 1
            continue
        # Tables: header, separator, data.
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, rows, refs)
            numbered_id = None
            bullet_id = None
            continue
        # Image.
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if m:
            path = (ROOT / m.group(2)).resolve()
            add_figure(doc, path, m.group(1))
            i += 1
            continue
        # Headings.
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            title = m.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            section_match = re.match(r"^(\d+)\.", title)
            if level == 1 and section_match and int(section_match.group(1)) in page_break_sections:
                p.paragraph_format.page_break_before = True
            if level == 2 and title == "Approval record":
                p.paragraph_format.page_break_before = True
            add_inline(p, title, refs)
            if level == 1:
                paragraph_border(p, CYAN, size=20, side="bottom")
            numbered_id = None
            bullet_id = None
            i += 1
            continue
        # Figure caption.
        if stripped.startswith("*Figure ") and stripped.endswith("*"):
            p = doc.add_paragraph(style="Figure Caption")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, stripped[1:-1], refs)
            i += 1
            continue
        # Blockquote callout.
        if stripped.startswith(">"):
            parts = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                parts.append(lines[i].strip().lstrip("> "))
                i += 1
            p = doc.add_paragraph(style="Callout")
            add_inline(p, " ".join(parts), refs)
            paragraph_border(p, MAGENTA, size=24, side="left")
            continue
        # Lists.
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            if bullet_id is None:
                bullet_id = add_numbering_definition(doc, ordered=False)
            p = doc.add_paragraph(style="YB List")
            apply_numbering(p, bullet_id)
            add_inline(p, m.group(1), refs)
            i += 1
            continue
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            if numbered_id is None:
                numbered_id = add_numbering_definition(doc, ordered=True, start=int(m.group(1)))
            p = doc.add_paragraph(style="YB List")
            apply_numbering(p, numbered_id)
            add_inline(p, m.group(2), refs)
            i += 1
            continue
        # Normal paragraph, joining wrapped source lines.
        parts = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt == "---"
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("![")
                or nxt.startswith(">")
                or nxt.startswith("```")
                or re.match(r"^[-*]\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)
                or re.match(r"^\[[^\]]+\]:\s+https?://", nxt)
            ):
                break
            parts.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(parts).replace("  ", " "), refs)
        numbered_id = None
        bullet_id = None


def build_docx():
    raw = SOURCE_MD.read_text(encoding="utf-8")
    lines = raw.splitlines()
    refs = parse_references(lines)
    # Body begins at numbered section 1; cover content is deliberately custom-built.
    body_start = next(i for i, line in enumerate(lines) if line.startswith("# 1. "))
    body = lines[body_start:]
    headings = []
    for line in body:
        m = re.match(r"^# (\d+)\.\s+(.*)$", line)
        if m:
            headings.append(m.group(2))

    doc = Document()
    configure_styles(doc)
    add_header_footer(doc.sections[0])
    props = doc.core_properties
    props.title = "YAWN.bot MVB-1 Engineering Specification"
    props.subject = "Buildable modular browser-card reference design"
    props.author = "YAWN.bot / yawn-ai"
    props.keywords = "YAWN.bot, Raspberry Pi Zero 2 W, browser kiosk, modular hardware, open source"
    props.comments = "Revision 1.0 — physical prototype validation required"
    add_cover(doc)
    add_contents(doc, headings)
    parse_body(doc, body, refs)

    # Ensure all sections use the same page and header/footer rules.
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


def main():
    BUILD.mkdir(parents=True, exist_ok=True)
    generate_figures()
    build_docx()
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
